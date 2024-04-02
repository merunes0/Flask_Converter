

import os
from flask import Flask, request, render_template
from docx import Document
from dotenv import load_dotenv

# Env variable
load_dotenv()
host = os.getenv("HOST")
port = os.getenv("PORT")

# Start
app = Flask(__name__, template_folder='templates')


@app.route('/')
def home():
    return render_template("converter.html")


@app.route('/uploader', methods=['POST'])
def upload_file():
    if 'file' in request.files:
        f = request.files['file']

        # File format check
        if not f.filename.endswith('.txt'):
            return 'Выбранный формат не является .txt форматом'

        file_content = f.read()
        file_decode = file_content.decode("utf-8")

        new_doc = Document()
        new_doc.add_paragraph(file_decode)

        # Encode Fix for Markdown (.md)
        lines = file_decode.split('\n')
        for line in lines:
            new_doc.add_paragraph(line)

        # File format

        if 'format' in request.values:
            selected_format = request.values['format']
            if selected_format == 'docx':
                file_path = f.filename.replace(".txt", ".docx")
                new_doc.save(file_path)
            elif selected_format == 'md':
                file_path = f.filename.replace(".txt", ".md")
                with open(file_path, 'w') as md_file:
                    for line in lines:
                        md_file.write(line + '\n')

            else:
                return 'Неверно выбран формат конвертации'

            return render_template("result.html", original_filename=f.filename, new_filename=file_path)

    return 'Error: File not selected or unsupported format'


if __name__ == '__main__':
    context = (host, port)
    app.run(debug=True, port=port, host=host)
